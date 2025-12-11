#!/usr/bin/env python3
"""Test script to demonstrate enhanced DOCX formatting with BSP tables"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Create a sample BSP examination report with tables
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Title
title_p = doc.add_paragraph()
title_run = title_p.add_run('SAMPLE BSP EXAMINATION REPORT')
title_run.font.size = Pt(16)
title_run.font.bold = True
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

# Section 1: Overall Summary
heading = doc.add_paragraph()
heading_run = heading.add_run('I. OVERALL SUMMARY OF ASSESSMENT')
heading_run.font.size = Pt(12)
heading_run.font.bold = True

doc.add_paragraph()

# Create assessment summary table
table1 = doc.add_table(rows=4, cols=2)
table1.style = 'Light Grid Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.LEFT

# Set column widths
table1.columns[0].width = Inches(4.0)
table1.columns[1].width = Inches(2.0)

# Header row
header_cells = table1.rows[0].cells
header_cells[0].text = 'Risk Category'
header_cells[1].text = 'Net Rating'

for cell in header_cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    set_cell_shading(cell, 'D9D9D9')

# Data rows
data = [
    ('Credit Risk', 'Moderate'),
    ('Liquidity Risk', 'Low'),
    ('Interest Rate Risk in the Banking Book (IRRBB)', 'Low')
]

for i, (category, rating) in enumerate(data, start=1):
    row = table1.rows[i]
    row.cells[0].text = category
    row.cells[1].text = rating
    
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if rating in paragraph.text:
                    run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()

# Section 2: Institutional Support
heading2 = doc.add_paragraph()
heading2_run = heading2.add_run('II. INSTITUTIONAL LEVEL SUPPORT (RATINGS)')
heading2_run.font.size = Pt(12)
heading2_run.font.bold = True

doc.add_paragraph()

# Create support ratings table
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.LEFT

# Set column widths
table2.columns[0].width = Inches(3.5)
table2.columns[1].width = Inches(2.5)

# Header row
header_cells2 = table2.rows[0].cells
header_cells2[0].text = 'Component'
header_cells2[1].text = 'Rating'

for cell in header_cells2:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
    set_cell_shading(cell, 'D9D9D9')

# Data rows
support_data = [
    ('Earnings', 'Strong'),
    ('Capital', 'Strong'),
    ('Liquidity', 'Strong'),
    ('Governance', 'Acceptable')
]

for i, (component, rating) in enumerate(support_data, start=1):
    row = table2.rows[i]
    row.cells[0].text = component
    row.cells[1].text = rating
    
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if rating in paragraph.text:
                    run.font.bold = True

doc.add_paragraph()

# Add sample narrative content
doc.add_paragraph()
narrative_heading = doc.add_paragraph()
narrative_heading_run = narrative_heading.add_run('III. ASSESSMENT NARRATIVE')
narrative_heading_run.font.size = Pt(12)
narrative_heading_run.font.bold = True

doc.add_paragraph()

narrative = doc.add_paragraph()
narrative.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
narrative_text = narrative.add_run(
    'The examination conducted revealed that the financial institution maintains adequate '
    'risk management practices across all assessed areas. The credit risk profile is rated '
    'as Moderate, reflecting a well-diversified loan portfolio with appropriate provisioning. '
    'Liquidity and interest rate risks are rated as Low, indicating robust treasury management '
    'and hedging strategies. The institution demonstrates Strong performance in earnings, capital, '
    'and liquidity metrics, with Acceptable governance frameworks in place.'
)
narrative_text.font.size = Pt(11)
narrative_text.font.name = 'Calibri'

# Save document
output_path = '/workspaces/fss-roe-writer/sample_bsp_format.docx'
doc.save(output_path)
print(f"✓ Sample BSP-formatted DOCX created: {output_path}")
